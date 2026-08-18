"""
Document Processor Module.
Responsible for reading multi-format document files (PDF, DOCX, CSV, TXT, MD, JSON),
extracting structured text content, and splitting documents into overlapping semantic
text chunks for embedding generation and vector retrieval.
"""

import os
import re
import csv
from typing import List, Dict, Any
from pypdf import PdfReader
import docx

try:
    import olefile
except ImportError:
    olefile = None

from app.config import CHUNK_SIZE, CHUNK_OVERLAP


def clean_text_content(text: str) -> str:
    """
    Purpose:
    Cleans up excessive whitespace, null characters, and unprintable characters.

    Parameters:
    text (str): Raw string input to sanitize.

    Returns:
    str: Sanitized clean string output.
    """
    if not text:
        return ""
    text = text.replace('\x00', '')
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    return re.sub(r'\s+', ' ', text).strip()


def is_zip_file(file_path: str) -> bool:
    """
    Purpose:
    Checks if a file starts with PK zip magic bytes (standard for OOXML .docx files).

    Parameters:
    file_path (str): Path to the target file.

    Returns:
    bool: True if ZIP header exists, otherwise False.
    """
    try:
        with open(file_path, 'rb') as f:
            header = f.read(4)
            return header.startswith(b'PK\x03\x04')
    except Exception:
        return False


def _extract_binary_doc_text(doc_path: str) -> List[str]:
    """
    Purpose:
    Extracts text strings from legacy binary MS Word (.doc) compound files.

    Parameters:
    doc_path (str): Path to the .doc file.

    Returns:
    List[str]: List of extracted paragraph strings.

    Raises:
    ValueError: If the file byte stream cannot be parsed.
    """
    paragraphs_text = []
    raw_bytes = b""

    # Attempt stream extraction via olefile if present and valid
    if olefile and olefile.isOleFile(doc_path):
        try:
            ole = olefile.OleFileIO(doc_path)
            streams = []
            for name in ['WordDocument', '1Table', '0Table']:
                if ole.exists(name):
                    streams.append(ole.openstream(name).read())
            if streams:
                raw_bytes = b'\n'.join(streams)
            ole.close()
        except Exception:
            raw_bytes = b""

    # Direct binary file read fallback
    if not raw_bytes:
        try:
            with open(doc_path, 'rb') as f:
                raw_bytes = f.read()
        except Exception as e:
            raise ValueError(f"Unable to read file bytes: {str(e)}")

    # Extract printable ASCII (min 5 chars) and UTF-16LE strings (min 5 chars)
    ascii_strings = [s.decode('ascii', errors='ignore') for s in re.findall(rb'[\x20-\x7e\x0a\x0d]{5,}', raw_bytes)]
    utf16_strings = [s.decode('utf-16le', errors='ignore') for s in re.findall(rb'(?:[\x20-\x7e\x0a\x0d]\x00){5,}', raw_bytes)]

    ignored_metadata = {
        'Root Entry', 'WordDocument', 'Table', 'SummaryInformation', 
        'DocumentSummaryInformation', 'CompObj', 'ObjectPool', 'Normal',
        'Default Paragraph Font', 'Header', 'Footer', 'Times New Roman',
        'Arial', 'Calibri', 'Courier New', 'Symbol', 'Wingdings'
    }

    for s in ascii_strings + utf16_strings:
        s_clean = clean_text_content(s)
        if len(s_clean) >= 10 and not any(s_clean.startswith(meta) for meta in ignored_metadata):
            paragraphs_text.append(s_clean)

    return paragraphs_text


def extract_text_from_pdf(pdf_path: str) -> List[Dict[str, Any]]:
    """
    Purpose:
    Extracts raw text content page-by-page from a PDF file.

    Parameters:
    pdf_path (str): Path to the PDF file.

    Returns:
    List[Dict[str, Any]]: List of pages containing 'page' index and 'text'.

    Raises:
    ValueError: If text extraction fails or no text is found.
    """
    try:
        reader = PdfReader(pdf_path)
        pages_content = []

        for index, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            cleaned = clean_text_content(page_text)
            if cleaned:
                pages_content.append({
                    "page": index + 1,
                    "text": cleaned
                })

        if not pages_content:
            raise ValueError("No extractable text found in the uploaded PDF.")

        return pages_content

    except Exception as e:
        raise ValueError(f"Failed to process PDF file '{pdf_path}': {str(e)}")


def extract_text_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """
    Purpose:
    Extracts text paragraphs and table contents from Word (.docx / .doc) documents.
    Handles OOXML (.docx) via python-docx, with a robust fallback for legacy binary Word (.doc) files.

    Parameters:
    docx_path (str): Path to the Word document.

    Returns:
    List[Dict[str, Any]]: Virtual pages of text containing 'page' index and 'text'.

    Raises:
    ValueError: If text extraction fails or no text is found.
    """
    paragraphs_text = []
    ext = os.path.splitext(docx_path)[1].lower()

    # Method 1: If it's a ZIP archive (standard for .docx), try python-docx first
    if ext == ".docx" or is_zip_file(docx_path):
        try:
            doc = docx.Document(docx_path)

            # Extract text from paragraphs
            for p in doc.paragraphs:
                txt = clean_text_content(p.text)
                if txt:
                    paragraphs_text.append(txt)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [clean_text_content(cell.text) for cell in row.cells if clean_text_content(cell.text)]
                    if row_cells:
                        paragraphs_text.append(" | ".join(row_cells))

        except Exception:
            # Fall back to binary decoding if python-docx fails
            paragraphs_text = []

    # Method 2: If python-docx returned nothing, or if file is legacy binary (.doc)
    if not paragraphs_text:
        try:
            paragraphs_text = _extract_binary_doc_text(docx_path)
        except Exception as fallback_err:
            raise ValueError(f"Failed to process Word document '{docx_path}': {str(fallback_err)}")

    if not paragraphs_text:
        raise ValueError(f"No extractable text found in Word document '{docx_path}'.")

    # Group paragraphs into virtual pages (~15 paragraphs or ~2500 characters per page)
    pages_content = []
    current_page_lines = []
    current_char_count = 0
    page_num = 1

    for line in paragraphs_text:
        current_page_lines.append(line)
        current_char_count += len(line)

        if current_char_count >= 2500 or len(current_page_lines) >= 15:
            pages_content.append({
                "page": page_num,
                "text": "\n".join(current_page_lines)
            })
            page_num += 1
            current_page_lines = []
            current_char_count = 0

    if current_page_lines:
        pages_content.append({
            "page": page_num,
            "text": "\n".join(current_page_lines)
        })

    return pages_content


def extract_text_from_csv(csv_path: str) -> List[Dict[str, Any]]:
    """
    Purpose:
    Extracts row-by-row key-value text from CSV spreadsheet documents.
    Formats each row with column headers for accurate semantic search.

    Parameters:
    csv_path (str): Path to the CSV file.

    Returns:
    List[Dict[str, Any]]: Virtual pages of text containing 'page' index and 'text'.

    Raises:
    ValueError: If text extraction fails or no text is found.
    """
    try:
        pages_content = []
        rows_per_page = 30
        
        # Detect encoding
        encoding = 'utf-8'
        try:
            with open(csv_path, 'r', encoding='utf-8') as f:
                f.read(2048)
        except UnicodeDecodeError:
            encoding = 'latin-1'

        with open(csv_path, 'r', encoding=encoding, errors='replace') as f:
            sample = f.read(4096)
            f.seek(0)
            
            # Determine if header exists
            has_header = True
            try:
                sniffer = csv.Sniffer()
                has_header = sniffer.has_header(sample)
            except Exception:
                has_header = True

            reader = csv.reader(f)
            headers = []
            formatted_rows = []

            for row_idx, row in enumerate(reader):
                if not row or not any(field.strip() for field in row):
                    continue

                if row_idx == 0 and has_header:
                    headers = [clean_text_content(h) or f"Column_{i+1}" for i, h in enumerate(row)]
                    continue

                if not headers:
                    headers = [f"Column_{i+1}" for i in range(len(row))]

                # Build row string "Header1: Val1 | Header2: Val2 ..."
                row_parts = []
                for i, val in enumerate(row):
                    clean_val = clean_text_content(val)
                    if clean_val:
                        col_name = headers[i] if i < len(headers) else f"Column_{i+1}"
                        row_parts.append(f"{col_name}: {clean_val}")

                if row_parts:
                    formatted_rows.append(f"Row {len(formatted_rows) + 1}: " + " | ".join(row_parts))

        if not formatted_rows:
            raise ValueError("No extractable data rows found in the uploaded CSV file.")

        # Batch rows into virtual pages
        page_num = 1
        for i in range(0, len(formatted_rows), rows_per_page):
            batch = formatted_rows[i:i + rows_per_page]
            pages_content.append({
                "page": page_num,
                "text": "\n".join(batch)
            })
            page_num += 1

        return pages_content

    except Exception as e:
        raise ValueError(f"Failed to process CSV file '{csv_path}': {str(e)}")


def extract_text_from_txt(file_path: str) -> List[Dict[str, Any]]:
    """
    Purpose:
    Extracts text from plain text, Markdown, JSON, or log files.

    Parameters:
    file_path (str): Path to the text/markdown/json document.

    Returns:
    List[Dict[str, Any]]: List of virtual pages containing 'page' index and 'text'.

    Raises:
    ValueError: If processing fails or text content is empty.
    """
    try:
        content = ""
        for enc in ['utf-8', 'latin-1', 'windows-1252']:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue

        cleaned = clean_text_content(content)
        if not cleaned:
            raise ValueError("The uploaded text file is empty.")

        # Split into blocks of ~2000 characters for page organization
        pages_content = []
        paragraphs = content.split('\n\n')
        current_page_text = []
        current_char_count = 0
        page_num = 1

        for p in paragraphs:
            clean_p = clean_text_content(p)
            if not clean_p:
                continue
            current_page_text.append(clean_p)
            current_char_count += len(clean_p)

            if current_char_count >= 2000:
                pages_content.append({
                    "page": page_num,
                    "text": "\n\n".join(current_page_text)
                })
                page_num += 1
                current_page_text = []
                current_char_count = 0

        if current_page_text:
            pages_content.append({
                "page": page_num,
                "text": "\n\n".join(current_page_text)
            })

        return pages_content

    except Exception as e:
        raise ValueError(f"Failed to process text document '{file_path}': {str(e)}")


def extract_text_from_file(file_path: str) -> List[Dict[str, Any]]:
    """
    Purpose:
    Unified entrypoint to extract text from files based on their extension.
    Supported formats: PDF, DOCX, DOC, CSV, TXT, MD, JSON, LOG.

    Parameters:
    file_path (str): Path to the document.

    Returns:
    List[Dict[str, Any]]: Extracted virtual page chunks of text content.

    Raises:
    ValueError: If the file type is unsupported or extraction fails.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext == ".csv":
        return extract_text_from_csv(file_path)
    elif ext in [".txt", ".md", ".json", ".log", ".text"]:
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format '{ext}'. Only PDF, Word, CSV, and Text files are supported.")


def chunk_document_pages(
    pages_content: List[Dict[str, Any]],
    doc_name: str,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP
) -> List[Dict[str, Any]]:
    """
    Purpose:
    Splits page-extracted text into overlapping chunks for dense vector embedding.
    Uses a sliding window approach with character boundaries while attempting
    to respect sentence boundaries to preserve semantic context.

    Parameters:
    pages_content (List[Dict[str, Any]]): Extracted text per page.
    doc_name (str): The original document filename for metadata tracking.
    chunk_size (int): Target character length per chunk. Default from config.
    chunk_overlap (int): Number of overlapping characters between adjacent chunks.

    Returns:
    List[Dict[str, Any]]: List of chunk dictionaries ready for embedding and indexing.
                          Each dict includes:
                          - 'id': Unique chunk string identifier
                          - 'text': Chunk text snippet
                          - 'metadata': File name, page number, chunk index
    """
    chunks = []
    global_chunk_idx = 0
    safe_doc_name = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', doc_name)

    for page_data in pages_content:
        page_num = page_data["page"]
        text = page_data["text"]

        if not text:
            continue

        start = 0
        text_length = len(text)

        while start < text_length:
            end = min(start + chunk_size, text_length)
            
            # If we're not at the end of text, try to snap to the nearest space or punctuation
            if end < text_length:
                boundary = max(text.rfind(' ', start, end), text.rfind('.', start, end))
                if boundary > start + (chunk_size // 2):
                    end = boundary + 1

            chunk_snippet = text[start:end].strip()

            if chunk_snippet:
                chunk_id = f"{safe_doc_name}_p{page_num}_c{global_chunk_idx}"
                chunks.append({
                    "id": chunk_id,
                    "text": chunk_snippet,
                    "metadata": {
                        "source": doc_name,
                        "page": page_num,
                        "chunk_index": global_chunk_idx,
                        "char_length": len(chunk_snippet)
                    }
                })
                global_chunk_idx += 1

            # Advance sliding window with step size (chunk_size - chunk_overlap)
            step = chunk_size - chunk_overlap
            if step <= 0:
                step = chunk_size
            start += step

    return chunks

